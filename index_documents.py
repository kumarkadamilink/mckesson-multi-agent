"""
index_documents.py
------------------
Reads all PDF and DOCX files from data/knowledge/contracts/,
chunks the text, embeds with fastembed (ONNX, no PyTorch needed),
and stores in a local ChromaDB collection.

Run from your project root:
    python index_documents.py

Requirements:
    pip install chromadb onnxruntime fastembed pypdf python-docx
"""

import re
import json
from pathlib import Path

import chromadb
from fastembed import TextEmbedding
from pypdf import PdfReader
from docx import Document

# ─────────────────────────────────────────────
# CONFIG — adjust paths if needed
# ─────────────────────────────────────────────
CONTRACTS_DIR   = Path("data/knowledge/contracts")
CHROMA_DIR      = Path("data/chroma_db")
COLLECTION_NAME = "procurement_contracts"
EMBED_MODEL     = "BAAI/bge-small-en-v1.5"   # fast ONNX model, 384-dim, no PyTorch

CHUNK_SIZE      = 400    # characters per chunk (smaller = less memory)
CHUNK_OVERLAP   = 50     # overlap between consecutive chunks
MAX_CHUNKS      = 50     # safety cap per file
BATCH_SIZE      = 10     # embed this many chunks at a time


# ─────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────
def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    elif suffix == ".docx":
        return extract_docx(path)
    else:
        print(f"  [SKIP] Unsupported file type: {path.name}")
        return ""


# ─────────────────────────────────────────────
# CHUNKING
# ─────────────────────────────────────────────
def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        # Always advance by at least (chunk_size - overlap) to guarantee progress
        start = start + chunk_size - overlap
    return chunks


# ─────────────────────────────────────────────
# METADATA EXTRACTION
# ─────────────────────────────────────────────
def extract_metadata(text: str, filename: str) -> dict:
    meta = {"filename": filename, "source": "contracts"}

    # Supplier name from filename e.g. Miller_Group_Contract.pdf -> Miller Group
    stem = Path(filename).stem
    stem = stem.replace("_Contract", "").replace("_", " ")
    meta["supplier"] = stem

    # Contract number
    m = re.search(r"Contract No[:\s]+([A-Z0-9\-]+)", text)
    if m:
        meta["contract_number"] = m.group(1).strip()

    # Effective date
    m = re.search(r"Effective[:\s]+([\d]{4}-[\d]{2}-[\d]{2})", text)
    if m:
        meta["effective_date"] = m.group(1)

    # Expiry date
    m = re.search(r"Expiry[:\s]+([\d]{4}-[\d]{2}-[\d]{2})", text)
    if m:
        meta["expiry_date"] = m.group(1)

    # Category
    categories = [
        "Medical Supplies", "Pharmaceutical", "IT Infrastructure",
        "Logistics", "Chemicals", "Facility Management",
        "Personal Protective Equipment", "Marketing", "Office Supplies",
        "Consulting"
    ]
    for cat in categories:
        if cat.lower() in text.lower():
            meta["category"] = cat
            break

    # Annual spend
    m = re.search(r"Total Annual Spend[^$]*\$([\d,]+)", text)
    if m:
        meta["annual_spend"] = "$" + m.group(1)

    # Payment terms
    m = re.search(r"Payment Terms\s*[|\s]+(Net \d+)", text)
    if m:
        meta["payment_terms"] = m.group(1)

    return meta


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def main():
    if not CONTRACTS_DIR.exists():
        print(f"ERROR: Directory not found: {CONTRACTS_DIR}")
        print("Make sure you are running from your project root.")
        return

    all_files = sorted(
        [f for f in CONTRACTS_DIR.iterdir()
         if f.suffix.lower() in (".pdf", ".docx")]
    )

    if not all_files:
        print(f"No PDF or DOCX files found in {CONTRACTS_DIR}")
        return

    # Deduplicate — prefer PDF, use DOCX only if no PDF twin exists
    stems_with_pdf = {f.stem for f in all_files if f.suffix.lower() == ".pdf"}
    contract_files = [
        f for f in all_files
        if not (f.suffix.lower() == ".docx" and f.stem in stems_with_pdf)
    ]

    print(f"Found {len(contract_files)} files to index "
          f"(PDFs preferred; DOCX used where no PDF exists)\n")

    # Load fastembed model
    # First run downloads ~130MB ONNX model and caches it locally
    print(f"Loading embedding model: {EMBED_MODEL}")
    print("  (First run downloads ~130MB — cached after that)\n")
    embed_model = TextEmbedding(model_name=EMBED_MODEL)

    # Init ChromaDB persistent client
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(path=str(CHROMA_DIR))

    # Drop existing collection for a clean reindex
    existing = [c.name for c in client.list_collections()]
    if COLLECTION_NAME in existing:
        print(f"Dropping existing collection '{COLLECTION_NAME}' for fresh index...")
        client.delete_collection(COLLECTION_NAME)

    collection = client.create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )
    print(f"Created collection: '{COLLECTION_NAME}'\n")

    total_chunks = 0
    index_log = []

    for file_path in contract_files:
        print(f"[{file_path.name}]")

        text = extract_text(file_path)
        if not text:
            print("  No text extracted — skipping.\n")
            continue

        print(f"  Extracted  : {len(text):,} characters")

        chunks = chunk_text(text)
        # Cap per file to avoid memory issues on large docs
        MAX_CHUNKS = 50
        if len(chunks) > MAX_CHUNKS:
            print(f"  Capping at {MAX_CHUNKS} chunks (was {len(chunks)})")
            chunks = chunks[:MAX_CHUNKS]
        print(f"  Chunks     : {len(chunks)}")

        meta = extract_metadata(text, file_path.name)
        print(f"  Supplier   : {meta.get('supplier')}")
        print(f"  Contract # : {meta.get('contract_number', 'n/a')}")
        print(f"  Spend      : {meta.get('annual_spend', 'n/a')}")
        print(f"  Category   : {meta.get('category', 'n/a')}")

        # Embed in batches of 8 to keep memory low
        BATCH_SIZE = 8
        embeddings = []
        for b in range(0, len(chunks), BATCH_SIZE):
            batch = chunks[b:b + BATCH_SIZE]
            embeddings.extend([e.tolist() for e in embed_model.embed(batch)])

        base_id = file_path.stem.lower().replace(" ", "_")
        ids = [f"{base_id}_chunk{i:03d}" for i in range(len(chunks))]
        metadatas = [
            {**meta, "chunk_index": i, "total_chunks": len(chunks)}
            for i in range(len(chunks))
        ]

        collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas,
        )

        total_chunks += len(chunks)
        index_log.append({
            "file": file_path.name,
            "characters": len(text),
            "chunks": len(chunks),
            "metadata": meta,
        })
        print()

    # Save index summary
    log_path = CHROMA_DIR / "index_log.json"
    with open(log_path, "w") as f:
        json.dump(index_log, f, indent=2)

    print("=" * 55)
    print(f"  Files indexed   : {len(index_log)}")
    print(f"  Total chunks    : {total_chunks}")
    print(f"  ChromaDB path   : {CHROMA_DIR}")
    print(f"  Collection      : {COLLECTION_NAME}")
    print(f"  Index log       : {log_path}")
    print("=" * 55)
    print("\nIndex ready. Run local_search_agent.py to query it.")


if __name__ == "__main__":
    main()
